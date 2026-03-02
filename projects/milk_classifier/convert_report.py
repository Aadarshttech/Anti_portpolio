import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            tag = 'w:{}'.format(side)
            element = tcPr.find(qn(tag))
            if element is not None:
                tcPr.remove(element)
            element = OxmlElement(tag)
            for key, val in kwargs[side].items():
                element.set(qn('w:{}'.format(key)), str(val))
            tcPr.append(element)

def convert_md_to_docx_advanced(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default Document Wide styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Adjust paragraph spacing
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by page separators
    pages = content.split('\n---\n')

    for i, page_content in enumerate(pages):
        lines = page_content.split('\n')
        in_code_block = False
        table_rows = []

        for line_raw in lines:
            line = line_raw.strip()

            # Handle Code Blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(line_raw) # Keep original spacing
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                continue

            # Handle Tables
            if '|' in line and '---' not in line:
                cells = [c.strip() for c in line.split('|') if c.strip() or (line.count('|') > 1)]
                if cells:
                    table_rows.append(cells)
                continue
            elif table_rows and (not line or '|' not in line):
                # End of table, render it
                if len(table_rows) > 0:
                    num_cols = max(len(row) for row in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    for r_idx, row_data in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < num_cols:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = cell_text
                                # Style cell text
                                for p in cell.paragraphs:
                                    p.paragraph_format.space_after = Pt(0)
                                    run = p.runs[0] if p.runs else p.add_run()
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                                    if r_idx == 0: run.bold = True
                table_rows = []
                if not line: continue

            # Skip table separators
            if '|' in line and '---' in line:
                continue

            # Main Chapter Headers (#)
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                continue

            # Sub Headers (##)
            if line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                continue

            # Sub-sub Headers (###)
            if line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                continue

            # Page Number markers (e.g., *Page ii*) - center them
            if line.startswith('*Page') and line.endswith('*'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[1:-1])
                run.font.italic = True
                run.font.size = Pt(10)
                continue

            # Centered text note (Centered on its own page) - skip or handle
            if '(Centered on its own page)' in line:
                continue

            # Bullet points
            if line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                text = line[2:]
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s*', '', line)
            else:
                p = doc.add_paragraph()
                text = line

            # Handle bold/italics in text
            if line:
                # Basic Bold Handling
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        # Add page break after each section except the last
        if i < len(pages) - 1:
            doc.add_page_break()

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

md_file = r"C:\Users\USER\.gemini\antigravity\brain\81fe8188-1c85-4ecb-a94d-8578c1e6f2a6\milk_quality_report.md"
docx_file = r"d:\Projects\Projects\Anti_portpolio\projects\milk_classifier\Milk_Quality_Report.docx"
convert_md_to_docx_advanced(md_file, docx_file)
